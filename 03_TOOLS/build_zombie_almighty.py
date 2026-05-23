from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(r"D:\Study\Project\zombie")
OUT_DIR = ROOT / "00_FINAL"
OUT_DOCX = OUT_DIR / "Zombie Almighty.docx"
OUT_PDF = OUT_DIR / "Zombie Almighty.pdf"


TITLE = "Zombie Almighty"
SUBTITLE = "Final SOT - SideCharacter model, latest hint fit, and clean test route"


CURRENT_ROUTE = """Gender reroll exactly once
-> zombify same-gender targets through passive Blood Bar + knock + manual grip
-> discover the real count during one clean run: check 40 / 45 / 49 / 50
-> if count alone does not pass, complete the tail:
   copy/mirror character appearance
   plus same-gender final zombify/check while the mirror relation is active"""


HARD_HINTS = [
    ("Gender Rerolling -> Zombifying -> ??? -> ???", "The main route is one reroll, then a same-gender zombification count, then two tail checks. The old repeated-reroll A/B/C matrix is deprecated."),
    ("Only one gender reroll", "Use exactly one gender reroll for the clean attempt. Extra gender rerolls are reset-risk and should not be part of the first route."),
    ("Same gender after reroll", "After the reroll, every count target should match the quest holder's post-reroll gender."),
    ("Amount is 50-ish; 57 is too much", "Run one staged count and check progress at 40, 45, 49, and 50. Do not waste time above 50 before tail probes."),
    ("Passive Blood Bar involved", "Build the target's Blood Bar through the Zombie passive interaction. Avoid T/G/Z/X/C as the clean Blood Bar source."),
    ("Manual gripping required", "Every conversion in the clean route must be knock -> manual grip -> confirmed zombie."),
    ("Zombies do not need to act", "After conversion, the target does nothing. No return, die, invade, chat, emote, kill, or movement requirement."),
    ("Can be done in base / A grade / any world", "No Volt, mode, raid, room, world, or high-grade transformation lock."),
    ("Any race; no other Quincy; no self-zombify", "Targets can be ordinary players of any race. The quest holder does not need to become a zombie."),
    ("Items/accessories are not needed", "If the tail uses appearance, treat it as character-copy/mirror state, not cosmetics, accessories, outfit, or item matching."),
    ("Same quest, different order / starting point", "The tail A/B order may vary per player. Swap only the tail order after the core count; do not return to broad reroll matrices."),
]


WHY_IT_FITS = [
    ("Any world / base", "The route only needs players, passive Blood Bar, manual grip, and NPC checks while on Zombie."),
    ("Zombification required", "The core work is repeated real player zombification."),
    ("Blood Bar passive", "Passive Blood Bar is the setup for every counted conversion."),
    ("Manual grip", "The route avoids move auto-zombify and uses manual grip for credit."),
    ("Gender swap required", "The one reroll is explicit and happens before the same-gender count."),
    ("Same-gender hint", "The count targets match the quest holder's post-reroll gender."),
    ("You do it twice / two tail tasks", "The visible route has count work plus two post-count checks; tail order can randomize."),
    ("TF2 autobalance us", "The gender axis is the balancing constraint; male/female path identity matters without requiring a physical formation."),
    ("You will not know what you did", "The exact count and tail direction can pass silently unless milestones and state are logged."),
    ("Become like them...then control their minds", "Become the required gender state, then create/control zombies through same-gender manual zombification."),
    ("No commands / zombie action", "The zombie exists as output; it does not perform the solution."),
    ("No accessories", "The appearance hypothesis is character-state mirroring, not item or cosplay matching."),
]


ROUTE_STEPS = [
    ("0", "Preparation", "Be on the Zombie Almighty step with Zombie equipped. Stay in base. Record current gender and server/world."),
    ("1", "One gender reroll", "Use exactly one gender reroll. The post-reroll gender becomes G for this run."),
    ("2", "Pick targets", "Use targets whose gender is G. Any race is allowed. Reuse is acceptable only if the server clearly counts it; fresh targets are cleaner for logs."),
    ("3", "Passive Blood Bar", "For each target, use the passive Blood Bar interaction. Do not use T/G/Z/X/C or Volt/mode as the clean setup."),
    ("4", "Manual zombify", "Knock the target, manually grip, and confirm the target becomes zombie. The zombie does nothing afterward."),
    ("5", "Milestone checks", "With Zombie still equipped, check Jugram/Balance only at 40, 45, 49, and 50. Do not talk while on the wrong Schrift."),
    ("6", "Tail primary", "If no pass by 50, copy/mirror another same-gender player's character appearance, then same-gender final zombify while that relation is active."),
    ("7", "Tail order swap", "If the primary tail fails, swap only the two tail checks: make/keep the mirror relation active before the final same-gender grip."),
    ("8", "Stop condition", "If both tail orders fail, stop and log. Do not add commands, mode, accessories, race locks, or extra gender rerolls."),
]


TAIL_DIRECTIONS = [
    ("1", "Quest holder copies/mirrors the same-gender target, then grips that target while the relation is active."),
    ("2", "Same-gender target copies/mirrors the quest holder, then gets gripped while the relation is active."),
    ("3", "Two same-gender helper targets mirror/copy each other, then the quest holder grips one while the relation is active."),
]


DO_NOT_DO = [
    "Do not run the old MMM/MMF/MFM/MFF/FFF/FFM/FMF/FMM A/B/C matrix as the main route.",
    "Do not run the old MM x2 -> FF x2 -> MF x2 -> FM x2 matrix.",
    "Do not gender reroll repeatedly during a clean attempt.",
    "Do not use return, die, invade, zombie commands, emotes, or chat strings.",
    "Do not use Volt/mode, Implode, or move-only Blood Bar setup for the clean route.",
    "Do not require Quincy, Soul Reaper, Arrancar, a specific Schrift, a specific world, or a clan room.",
    "Do not require the quest holder to be zombified.",
    "Do not require items, accessories, name, outfit, or cosplay matching.",
]


RUN_LOG = """Run ID:
Date:
Quest holder:
Zombie equipped before NPC checks: YES / NO
Initial gender:
Post-reroll gender G:
Gender rerolls used: exactly one / other
Targets gender: same as G / mixed
Target races:
Passive Blood Bar method used:
Manual grip confirmed for every target: YES / NO
Zombie confirmed for every target: YES / NO
NPC at 40:
NPC at 45:
NPC at 49:
NPC at 50:
Tail direction tested:
Mirror/copy relation active: YES / NO
Final same-gender grip after mirror/copy: YES / NO
Commands used: NONE
Volt/mode used: NONE
Wrong Schrift NPC talk: NONE
Final NPC result:
Pass/fail/blocked:
Notes:"""


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, *, bold: bool = False, fill: str | None = None, color: str = "1f2933", size: float = 8.7) -> None:
    if fill:
        shade(cell, fill)
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_heading(text, level=level)
    para.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    para.paragraph_format.space_after = Pt(3)
    for run in para.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor(24, 55, 50)


def add_para(doc: Document, text: str, *, bold: bool = False) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9.4)
    run.font.color.rgb = RGBColor(31, 41, 51)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(1)
        run = para.add_run(item)
        run.font.name = "Aptos"
        run.font.size = Pt(8.9)


def add_code(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, "F3F7F5")
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    for line in text.splitlines():
        run = para.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(8.2)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell(table.cell(0, i), header, bold=True, fill="183732", color="FFFFFF", size=8.8)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value, bold=(i == 0), size=8.5)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def build_docx() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(9.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor(24, 55, 50)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subtitle.add_run(SUBTITLE)
    srun.font.name = "Aptos"
    srun.font.size = Pt(10)
    srun.font.color.rgb = RGBColor(88, 101, 105)

    add_heading(doc, "1. Final Answer")
    add_para(doc, "Use this as the current final. Older A/B/C, 2x2, command, cosmetic, and Volt/mode theories are deprecated for clean testing.", bold=True)
    add_code(doc, CURRENT_ROUTE)

    add_heading(doc, "2. Latest Hard Hint Map")
    add_table(doc, ["Hint / constraint", "Final interpretation"], HARD_HINTS, [2.55, 4.75])

    add_heading(doc, "3. Clean Route")
    add_table(doc, ["Step", "Phase", "Action"], ROUTE_STEPS, [0.45, 1.55, 5.25])

    add_heading(doc, "4. Tail Direction Priority")
    add_para(doc, "If the count does not pass by 50, test the character-copy / mirror / gender tail before commands, mode routes, or old matrices.")
    add_table(doc, ["Priority", "Tail direction"], TAIL_DIRECTIONS, [0.75, 6.35])

    add_heading(doc, "5. Why This Satisfies The Hints")
    add_table(doc, ["Hint", "How the final route satisfies it"], WHY_IT_FITS, [2.15, 5.1])

    add_heading(doc, "6. Do Not Add These")
    add_bullets(doc, DO_NOT_DO)

    add_heading(doc, "7. Result Logging")
    add_code(doc, RUN_LOG)

    add_heading(doc, "8. Confidence")
    add_para(doc, "Confidence is high for the mechanism family: one reroll, same-gender passive Blood Bar, manual grip, no commands, no mode, any race.")
    add_para(doc, "Confidence is medium for the exact count and mirror direction until an in-game NPC pass is logged. The best immediate test is CORE_M_50 or CORE_F_50 depending on post-reroll gender, with milestone checks at 40/45/49/50.")

    doc.save(OUT_DOCX)
    return OUT_DOCX


def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("title", parent=base["Title"], fontName="Helvetica-Bold", fontSize=22, leading=26, alignment=TA_CENTER, textColor=colors.HexColor("#183732"), spaceAfter=4),
        "subtitle": ParagraphStyle("subtitle", parent=base["Normal"], fontName="Helvetica", fontSize=10, leading=14, alignment=TA_CENTER, textColor=colors.HexColor("#586569"), spaceAfter=10),
        "h1": ParagraphStyle("h1", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=13, leading=17, textColor=colors.HexColor("#183732"), spaceBefore=8, spaceAfter=4),
        "body": ParagraphStyle("body", parent=base["BodyText"], fontName="Helvetica", fontSize=8.7, leading=12, textColor=colors.HexColor("#1f2933"), spaceAfter=3),
        "code": ParagraphStyle("code", parent=base["Code"], fontName="Courier", fontSize=7.8, leading=10, backColor=colors.HexColor("#F3F7F5"), leftIndent=6, rightIndent=6, spaceBefore=4, spaceAfter=6),
        "cell": ParagraphStyle("cell", parent=base["BodyText"], fontName="Helvetica", fontSize=7.3, leading=9.2, textColor=colors.HexColor("#1f2933")),
        "cell_bold": ParagraphStyle("cell_bold", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=7.3, leading=9.2, textColor=colors.HexColor("#1f2933")),
    }


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), style)


def _pdf_table(headers: list[str], rows: list[tuple[str, ...]], widths_mm: list[float], styles: dict[str, ParagraphStyle]) -> Table:
    data = [[_p(h, styles["cell_bold"]) for h in headers]]
    for row in rows:
        data.append([_p(str(v), styles["cell_bold"] if i == 0 else styles["cell"]) for i, v in enumerate(row)])
    table = Table(data, colWidths=[w * mm for w in widths_mm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCEBE7")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#9FB7B0")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7FAF9")]),
            ]
        )
    )
    return table


def build_pdf() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    styles = _pdf_styles()
    doc = SimpleDocTemplate(str(OUT_PDF), pagesize=A4, leftMargin=16 * mm, rightMargin=16 * mm, topMargin=16 * mm, bottomMargin=16 * mm, title=TITLE)
    story = [
        Paragraph(TITLE, styles["title"]),
        Paragraph(SUBTITLE, styles["subtitle"]),
        Paragraph("1. Final Answer", styles["h1"]),
        _p("Use this as the current final. Older A/B/C, 2x2, command, cosmetic, and Volt/mode theories are deprecated for clean testing.", styles["body"]),
        Preformatted(CURRENT_ROUTE, styles["code"]),
        Paragraph("2. Latest Hard Hint Map", styles["h1"]),
        _pdf_table(["Hint / constraint", "Final interpretation"], HARD_HINTS, [52, 118], styles),
        Paragraph("3. Clean Route", styles["h1"]),
        _pdf_table(["Step", "Phase", "Action"], ROUTE_STEPS, [12, 34, 124], styles),
        Paragraph("4. Tail Direction Priority", styles["h1"]),
        _p("If the count does not pass by 50, test the character-copy / mirror / gender tail before commands, mode routes, or old matrices.", styles["body"]),
        _pdf_table(["Priority", "Tail direction"], TAIL_DIRECTIONS, [20, 150], styles),
        PageBreak(),
        Paragraph("5. Why This Satisfies The Hints", styles["h1"]),
        _pdf_table(["Hint", "How the final route satisfies it"], WHY_IT_FITS, [48, 122], styles),
        Paragraph("6. Do Not Add These", styles["h1"]),
    ]
    for item in DO_NOT_DO:
        story.append(_p("- " + item, styles["body"]))
    story.extend(
        [
            Paragraph("7. Result Logging", styles["h1"]),
            Preformatted(RUN_LOG, styles["code"]),
            Paragraph("8. Confidence", styles["h1"]),
            _p("Confidence is high for the mechanism family: one reroll, same-gender passive Blood Bar, manual grip, no commands, no mode, any race.", styles["body"]),
            _p("Confidence is medium for the exact count and mirror direction until an in-game NPC pass is logged. The best immediate test is CORE_M_50 or CORE_F_50 depending on post-reroll gender, with milestone checks at 40/45/49/50.", styles["body"]),
        ]
    )
    story = [x if isinstance(x, Spacer) else x for x in story]
    doc.build(story)
    return OUT_PDF


def build_all() -> tuple[Path, Path]:
    docx = build_docx()
    pdf = build_pdf()
    print(f"DOCX written to: {docx}")
    print(f"PDF written to:  {pdf}")
    return docx, pdf


if __name__ == "__main__":
    build_all()
